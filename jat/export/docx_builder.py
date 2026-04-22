"""Word (.docx) document builder for the Job Application Tracker."""

import sqlite3
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from jat.analytics import queries as analytics_queries


def _iso_to_display(date_str: str | None) -> str:
    """Convert YYYY-MM-DD to dd.mm.yyyy, or '---' when None."""
    if not date_str:
        return "---"
    try:
        return date.fromisoformat(date_str).strftime("%d.%m.%Y")
    except ValueError:
        return str(date_str)


def _format_salary_range(currency: str | None, min_val, max_val) -> str:
    """Return a formatted salary range string or '---' when all absent."""
    if currency is None and min_val is None and max_val is None:
        return "---"
    parts: list[str] = []
    if currency:
        parts.append(str(currency))
    if min_val is not None:
        parts.append(str(int(min_val)))
    if min_val is not None and max_val is not None:
        parts.append("--")
    if max_val is not None:
        parts.append(str(int(max_val)))
    return " ".join(parts) if parts else "---"


def _safe(value) -> str:
    """Return str(value) or empty string for None."""
    return str(value) if value is not None else ""


def _bold_row(row) -> None:
    """Make all runs in the first cell of each cell in a table row bold."""
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True


def _set_cell_text_bold(cell, text: str) -> None:
    """Set cell text and make it bold."""
    cell.text = text
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True


def build_afa_table(
    conn: sqlite3.Connection,
    output_path: str,
    date_from: str | None = None,
    date_to: str | None = None,
    status_id: int | None = None,
) -> None:
    """Build an AfA table Word document and save it to output_path."""
    conditions: list[str] = []
    params: list = []
    if date_from:
        conditions.append("a.application_date >= ?")
        params.append(date_from)
    if date_to:
        conditions.append("a.application_date <= ?")
        params.append(date_to)
    if status_id is not None:
        conditions.append("a.status_id = ?")
        params.append(status_id)

    where = ("WHERE " + " AND ".join(conditions)) if conditions else ""

    rows = conn.execute(
        f"""
        SELECT a.application_date, c.company_name, a.role_title,
               rsrc.label AS source_label, rs.label AS status_label, a.city
        FROM applications a
        LEFT JOIN companies c ON c.company_id = a.company_id
        LEFT JOIN ref_statuses rs ON rs.id = a.status_id
        LEFT JOIN ref_sources rsrc ON rsrc.id = a.source_id
        {where}
        ORDER BY a.application_date DESC
        """,
        params,
    ).fetchall()

    status_label = "Alle"
    if status_id is not None:
        status_row = conn.execute(
            "SELECT label FROM ref_statuses WHERE id = ?", (status_id,)
        ).fetchone()
        if status_row:
            status_label = _safe(status_row[0])

    doc = Document()
    doc.add_heading("Bewerbungsliste", level=1)

    from_str = _iso_to_display(date_from) if date_from else "Alle"
    to_str = _iso_to_display(date_to) if date_to else "Alle"
    today_str = date.today().strftime("%d.%m.%Y")
    doc.add_paragraph(
        f"Erstellt: {today_str}   Von: {from_str}   Bis: {to_str}   "
        f"Status: {status_label}   Gesamt: {len(rows)}"
    )

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Datum", "Unternehmen", "Stelle", "Quelle", "Status", "Ort"]):
        _set_cell_text_bold(hdr[i], h)

    for row in rows:
        cells = table.add_row().cells
        values = [
            _iso_to_display(row[0]),
            _safe(row[1]),
            _safe(row[2]),
            _safe(row[3]),
            _safe(row[4]),
            _safe(row[5]),
        ]
        for i, v in enumerate(values):
            cells[i].text = v

    doc.save(output_path)


def build_application_sheet(
    conn: sqlite3.Connection,
    output_path: str,
    application_id: int,
) -> None:
    """Build an application detail sheet Word document and save to output_path."""
    row = conn.execute(
        """
        SELECT
            a.role_title, c.company_name, a.application_date,
            a.response_date, rs.label AS status_label,
            rc.label AS category_label,
            ret.label AS employment_type_label,
            rsrc.label AS source_label,
            rw.label AS work_mode_label,
            a.city, a.country, rcur.label AS currency_label,
            a.salary_min, a.salary_max,
            a.priority_score, a.follow_up_date,
            a.job_posting_url, a.notes
        FROM applications a
        LEFT JOIN companies c ON c.company_id = a.company_id
        LEFT JOIN ref_statuses rs ON rs.id = a.status_id
        LEFT JOIN ref_categories rc ON rc.id = a.category_id
        LEFT JOIN ref_employment_types ret ON ret.id = a.employment_type_id
        LEFT JOIN ref_sources rsrc ON rsrc.id = a.source_id
        LEFT JOIN ref_work_modes rw ON rw.id = a.work_mode_id
        LEFT JOIN ref_currencies rcur ON rcur.id = a.currency_id
        WHERE a.application_id = ?
        """,
        (application_id,),
    ).fetchone()

    if row is None:
        raise ValueError(f"Application {application_id} not found.")

    salary = _format_salary_range(row[11], row[12], row[13])

    doc = Document()
    doc.add_heading(_safe(row[0]), level=1)
    doc.add_heading(_safe(row[1]), level=2)

    fields = [
        ("Bewerbungsdatum", _iso_to_display(row[2])),
        ("Antwortdatum", _iso_to_display(row[3])),
        ("Status", _safe(row[4])),
        ("Kategorie", _safe(row[5])),
        ("Anstellungsart", _safe(row[6])),
        ("Quelle", _safe(row[7])),
        ("Arbeitsmodell", _safe(row[8])),
        ("Stadt", _safe(row[9])),
        ("Land", _safe(row[10])),
        ("Gehalt", salary),
        ("Priorität", _safe(row[14])),
        ("Follow-up", _iso_to_display(row[15])),
        ("Stelle URL", _safe(row[16])),
        ("Notizen", _safe(row[17])),
    ]

    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in fields:
        cells = table.add_row().cells
        _set_cell_text_bold(cells[0], label)
        cells[1].text = value

    doc.save(output_path)


def build_analytics_summary(
    conn: sqlite3.Connection,
    output_path: str,
    date_from: str | None = None,
    date_to: str | None = None,
) -> None:
    """Build an analytics summary Word document and save it to output_path."""
    stats = analytics_queries.summary_stats(conn)
    status_rows = analytics_queries.applications_by_status(conn)
    total_for_pct = sum(r["count"] for r in status_rows) or 1

    doc = Document()
    doc.add_heading("Analytik-Bericht", level=1)

    from_str = _iso_to_display(date_from) if date_from else "Alle"
    to_str = _iso_to_display(date_to) if date_to else "Alle"
    today_str = date.today().strftime("%d.%m.%Y")
    doc.add_paragraph(f"Erstellt: {today_str}   Von: {from_str}   Bis: {to_str}")

    doc.add_heading("Zusammenfassung", level=2)
    response_pct = f"{stats['response_rate'] * 100:.0f}%"
    summary_fields = [
        ("Bewerbungen gesamt", str(stats["total_applications"])),
        ("Aktive Bewerbungen", str(stats["active_applications"])),
        ("Antwortrate", response_pct),
        ("Durchschn. Priorität", str(stats["avg_priority"])),
    ]
    stats_table = doc.add_table(rows=0, cols=2)
    stats_table.style = "Table Grid"
    for label, value in summary_fields:
        cells = stats_table.add_row().cells
        _set_cell_text_bold(cells[0], label)
        cells[1].text = value

    doc.add_heading("Status-Verteilung", level=2)
    status_table = doc.add_table(rows=1, cols=3)
    status_table.style = "Table Grid"
    hdr = status_table.rows[0].cells
    for i, h in enumerate(["Status", "Anzahl", "Anteil"]):
        _set_cell_text_bold(hdr[i], h)

    for r in status_rows:
        pct = f"{r['count'] / total_for_pct * 100:.0f}%"
        cells = status_table.add_row().cells
        cells[0].text = _safe(r["label"])
        cells[1].text = str(r["count"])
        cells[2].text = pct

    doc.save(output_path)
